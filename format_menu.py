# https://pypi.org/project/xlrd/
# https://pypi.org/project/python-docx/

import sys
import xlrd
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----- Settings -----

# Arbitrary number, it is not the width of a row.
row_width = 39

display_dots = True

# Section's font style, font size, bold, italicize
section_font = 'Chow Fun'
section_font_size = 10.5
section_bold = True
section_italic = False

# Dish name's font style, font size, bold, italicize
dish_font = 'Hack Bold'
dish_font_size = 8
dish_bold = False
dish_italic = False

# Description's font style, font size, bold, italicize
description_font = 'Calibri'
description_font_size = 8
description_bold = False
description_italic = True

input_excel_filename = 'menu'
output_filename = 'staging'

# 1 = 1, 2 = 2, ...
# Column names is on row 1 so start row is 2
start_row = 2
end_row = 5

# A = 1, B = 2, ...
section_column = 2
dish_name_column = 1
price_column = 3
description_column = 4

# Character to repeat
character = "."

# ----- End of Settings -----


def forced_text_wrap(text, max_length):
    """
    Returns a list of the text broken into multiple lines constrained by
    max_length on each line.
    """
    if (len(text) < max_length):
        return [text]

    lines = []
    tmp = text

    while True:
        # If text is longer than max_length
        if (len(tmp) > max_length):
            # Break beginning chunk at max_length and append to list
            break_index = tmp.rfind(" ", 0, max_length+1)

            # Error check if text cannot break into smaller chunks on spaces
            if break_index == -1:
                ex = 'Cannot break following text into multiple lines based on'
                'the constraint:\n' + str(max_length) + '\n' + text
                raise Exception(ex)

            lines.append(tmp[:break_index])

            # Assign tmp as the rest of text
            tmp = tmp[break_index+1:]
        else:
            # Else text is smaller or equal max_length,
            # append rest of text to list and break loop
            lines.append(tmp)
            break

    return lines


def write_row(dish_name, price, description):
    """
    Write dish_name, price, and description into document.
    """
    # Write dish_name
    # Split dish_name string into multiple lines constrained by row_width
    lines = forced_text_wrap(dish_name, row_width)

    # Write lines into document
    for x in range(0, len(lines)):
        p = document.add_paragraph(lines[x], style=dish_style)

    # Append a space at the end
    p.add_run(" ")

    # Append x number of dots/spacer to progress 'cursor' to a certain column
    spacer_size = row_width - len(lines[len(lines)-1])
    for x in range(0, spacer_size):
        if display_dots:
            p.add_run(character)
        else:
            p.add_run(" ")

    # Write price
    price_ = '%.2f' % (price)
    if len(price_) == 4:
        if display_dots:
            p.add_run(character)
        else:
            p.add_run(" ")
    p.add_run(" " + price_)

    # Write description
    if description != xlrd.empty_cell.value:
        tmp = description.split("\\n")
        for x in range(0, len(tmp)):
            document.add_paragraph(tmp[x], style=description_style)
        #document.add_paragraph(description, style=description_style)


if len(sys.argv) > 1:
    row_width = int(sys.argv[1])

document = Document()

styles = document.styles

dish_style = "dish_style"
description_style = "description_style"
section_style = "section_style"

style = styles.add_style(dish_style, WD_STYLE_TYPE.PARAGRAPH)
style.font.name = dish_font
style.font.size = Pt(dish_font_size)
style.font.bold = dish_bold
style.font.italic = dish_italic

style = styles.add_style(description_style, WD_STYLE_TYPE.PARAGRAPH)
style.font.name = description_font
style.font.size = Pt(description_font_size)
style.font.bold = description_bold
style.font.italic = description_italic

style = styles.add_style(section_style, WD_STYLE_TYPE.PARAGRAPH)
style.font.name = section_font
style.font.size = Pt(section_font_size)
style.font.bold = section_bold
style.font.italic = section_italic

workbook = xlrd.open_workbook(input_excel_filename + '.xlsx')
sheet = workbook.sheet_by_index(0)

current_section = ""

for row in range(start_row-1, end_row):
    section = sheet.cell(row, section_column-1).value
    dish_name = sheet.cell(row, dish_name_column-1).value
    price = sheet.cell(row, price_column-1).value
    description = sheet.cell(row, description_column-1).value

    if section != current_section:
        # Write section
        current_section = section
        document.add_paragraph(
            "\n" + section,
            style=section_style
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Write row (dish + price + description)
    write_row(dish_name, price, description)

# Save document as microsoft word docx
document.save(output_filename + '.docx')
